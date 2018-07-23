from docx import *
document = Document('abc.docx')

a=1
#read paragraph 
for par in document.paragraphs: 
    pa=par.text
    pri = str(a)+" : "+str(pa)
    print(pri)
    a=a+1

b=1
#read table 
for table in document.tables:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                pri = str(b)+" : "+str(paragraph.text)
                print(pri)
                b=b+1

print()
#Author 
print(document.core_properties.author)

#Title
print(document.core_properties.title)

#Category
print(document.core_properties.category)

#Language
print(document.core_properties.language)

#Subject
print(document.core_properties.subject)

"""
Code to extract data in xml format
body_element = document._body._body
print(body_element.xml) 

"""
